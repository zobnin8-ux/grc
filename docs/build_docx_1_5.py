# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def title(text):
    doc.add_heading(text, level=0)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            t.rows[ri + 1].cells[ci].text = cell
    doc.add_paragraph()

title("Американский сайт: стратегия (пункты 1–5)")
p("Краткая версия · май 2026")
p("Полный документ: Мастер-документ-сайт-США.md")
p("Конкурент: tissinc.com · Внутренний ориентир capability: 1grc.ru (без России на US-сайте)")

h1("1. Контекст проекта")
h2("1.1 Задача")
p("Создать современный американский многостраничный сайт для industrial-компании на рынке Houston / Gulf Coast (mobilization по США — в разумных пределах).")
p("Сайт должен:")
bullets([
    "НЕ выглядеть как перевод русского сайта",
    "НЕ выглядеть как стартап, AI-компания или SaaS-лендинг",
    "НЕ выглядеть как «маленький мобильный ремонт в фургоне»",
    "ДА выглядеть как full-scale Gulf Coast industrial field operations company с rapid deployment",
    "ДА превосходить старых локальных конкурентов (TISS и аналоги) по trust, CTA и первому экрану",
])

h2("1.2 Бизнес-модель (зафиксировано)")
table(
    ["Параметр", "Решение"],
    [
        ("Россия", "Бизнес полностью в США. На сайте нет России, RU, российских заказчиков"),
        ("Юрлицо", "US LLC (имя, домен, адрес, телефон, email — placeholders)"),
        ("Партнёры", "Не упоминать"),
        ("Люди", "Все в Америке"),
        ("Оборудование", "В США (целевая модель)"),
        ("Язык", "Только English"),
        ("Capability", "Факты с 1grc.ru → US-структура, EN copy"),
        ("Кейсы на launch", "~6 на EN, без имён российских заводов"),
        ("География", "Houston-based, Gulf Coast + честная строка о wider mobilization"),
        ("Фото", "Сток на старте → замена от клиента"),
        ("Insurance / certs", "После LLC/полиса; на launch — гибкий Trust block"),
        ("Формы (MVP)", "Email"),
    ],
)

h2("1.3 Принцип структуры")
p("Не 13 равных страниц, а: 10 сильных + 2 коротких + 1 hub (карта услуг).", bold=True)
p("Полный спектр — в меню; глубина — field machining, emergency, shutdown, projects, equipment.")

h1("2. Анализ конкурента TISS")
p("Сайт: https://tissinc.com/")
p("Профиль: millwright, field machining, shop machining · refining, petrochemical, power, midstream, offshore · Houston (281.922.7555).")

h2("2.1 Оценка «Сайт Доктор»")
table(
    ["Метрика", "Балл / 100"],
    [
        ("Общий", "48"),
        ("Доверие", "42"),
        ("CTA", "25 — критическая слабость"),
        ("Ясность", "55"),
        ("Mobile", "62"),
    ],
)
p("Плюсы: структура, три услуги, телефон в шапке, нормальное меню.")
p("Минусы:")
bullets([
    "На главной нет CTA (форма только на Contact; номер без tel:)",
    "Нет trust на главной — шестерёнки, без фото/кейсов/цифр",
    "Заголовок — leading provider… без УТП",
    "Нет H1, нет meta description",
    "Главная = один экран; flip-карточки с пустой обратной стороной",
    "Сильные цифры (150\" VTL, 40' lathe) спрятаны в Shop",
    "Сайт устарел (ощущение brochure ~2014)",
])

h2("2.2 Вердикт")
p("TISS — старая визитка, не conversion-ready industrial contractor.", bold=True)
table(
    ["Критерий", "TISS сейчас"],
    [
        ("Первые 5 секунд", "Отрасли ясны, «почему мы» — нет"),
        ("Proof", "Почти ноль на главной"),
        ("Leads", "Форма спрятана"),
    ],
)

h2("2.3 Эталон заголовка (Site Doctor)")
p("Было: …leading provider of fixed and rotating machinery services to the refining, petrochemical…")
p("Надо (логика): Industrial equipment service for petrochemical and energy — from field repairs to lathe work with diameters up to 150 inches.")

h2("2.4 Где бить TISS")
bullets([
    "Hero + CTA за 5 секунд",
    "Projects, Equipment, фото, цифры",
    "Формы + emergency; tel:; mobile sticky CTA",
    "H1, meta, operational dark UI",
    "Rapid deployment (карта, urgency)",
])
p("Цель: выиграть сравнение за 30 секунд, не SEO домена за 3 месяца.", bold=True)

h1("3. Сравнение с 1grc.ru — уроки")
p("Сайт: https://www.1grc.ru/ · Site Doctor: 48/100 (ясность 38, SEO 37).")

h2("3.1 Не повторять на US")
bullets([
    "Размытый hero («адаптивное производство»)",
    "H1 не совпадает с первым экраном",
    "Доверие и кейсы ниже fold, простыня отзывов",
    "CTA «обратный звонок» без обещания",
    "Несколько брендов на одном сайте",
])

h2("3.2 Перенести (перепаковать)")
bullets([
    "Мобильная обработка на объекте, свои станки",
    "Кейсы: задача → метод → срок",
    "Отрасли → US wording",
    "Сильный title → в US hero + H1",
])

h2("3.3 Зеркало")
table(
    ["", "1grc.ru", "TISS", "Наш US сайт"],
    [
        ("Балл SD", "48", "48", "Выигрыш по конверсии"),
        ("Слабое", "Витрина", "Главная пустая", "—"),
        ("Сильное", "Контент", "Ясность меню", "Витрина + proof"),
    ],
)

h1("4. Что значит «убить конкурентов»")
p("Реалистично:", bold=True)
bullets([
    "Рядом с TISS наш сайт = operationally ready, их = old brochure",
    "Заявка менее чем за 2 клика на mobile",
    "COI / MSA / ISNet — после LLC, не выдумывать на launch",
])
p("Нереалистично обещать на сайте:", bold=True)
bullets([
    "SEO #1 за месяц",
    "Millwright «всё» без подтверждённого capability",
    "24/7 nationwide за 2 часа без диспетчеризации",
])

h1("5. Стратегическое позиционирование")
h2("5.1 Главная идея")
p("НЕ: We are a mobile repair company.")
p("ДА: We are a full-scale industrial field operations company with rapid deployment capability.", bold=True)
p("Mobile crews — преимущество, не весь бренд.")

h2("5.2 Целевая реакция")
p("These people look operationally ready.")

h2("5.3 Field-first")
table(
    ["Уровень", "Содержание"],
    [
        ("Hero", "Field machining, on-site repair, emergency, shutdown, mobile crews"),
        ("Proof", "Projects, Equipment, process, trust (гибкий)"),
        ("Меню", "Полный спектр на hub; Rotating/Mechanical — коротко или v2"),
        ("Shop", "Поддержка, когда есть US yard"),
    ],
)

h2("5.4 Millwright vs field machining")
bullets([
    "Field machining — ядро (металлообработка на объекте)",
    "Millwright — только если реально в US; не клонировать меню TISS без scope",
    "Launch: Field + Emergency + Projects + Equipment — глубже остального",
])

h2("5.5 Credibility (без России)")
bullets([
    "Houston-based. Now operating from Houston.",
    "Кейсы без Severstal, Mondelēz и т.д.",
    "Факты с 1grc.ru → рерайт, не дословный перевод",
    "About (опционально): decades of in-field machining and proprietary mobile equipment — без 1grc.ru в публичном поле",
])

p("Пункты 6–18 (структура страниц, дизайн, техника, формы) — в файле Мастер-документ-сайт-США.md.")

out = r"D:\ArtemSite\docs\Strategiya-sait-USA-punkty-1-5.docx"
doc.save(out)
